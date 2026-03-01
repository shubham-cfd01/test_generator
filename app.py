import os
import json
import traceback
import io as io_module
import math
import random
import re as _re
import sqlite3
import hashlib
import functools
import pandas as pd
from flask import Flask, render_template, request, send_file, jsonify, Response, redirect, url_for, session
from groq import Groq
from io import BytesIO
import threading
import uuid
import time
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.utils import simpleSplit, ImageReader
from reportlab.pdfgen import canvas
import matplotlib
matplotlib.use('Agg')  # non-interactive backend
import matplotlib.pyplot as plt

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'arth-academy-secret-key-change-in-prod')

ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'arth@admin2026')
# Use persistent path on Railway (set DATABASE_PATH=/data/users.db and add Volume at /data)
# Otherwise use local data/ for development - data is lost on redeploy without a volume
_vol = os.environ.get('RAILWAY_VOLUME_MOUNT_PATH', '')
_db_env = os.environ.get('DATABASE_PATH', '')
if _db_env:
    DB_PATH = _db_env
elif _vol:
    DB_PATH = os.path.join(_vol, 'users.db')
else:
    DB_PATH = os.path.join(os.path.dirname(__file__), 'data', 'users.db')

# ─────────────────────────────────────────────────────────
# Global memory store for progress tracking
# ─────────────────────────────────────────────────────────
progress_store = {}
result_store = {}


# ─────────────────────────────────────────────────────────
# USER ACCESS DATABASE
# ─────────────────────────────────────────────────────────
def init_db():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.execute('''CREATE TABLE IF NOT EXISTS allowed_users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        contact TEXT UNIQUE NOT NULL,
        password TEXT NOT NULL DEFAULT '',
        name TEXT DEFAULT '',
        products TEXT NOT NULL DEFAULT 'both',
        added_on TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        is_active INTEGER DEFAULT 1
    )''')
    for col, default in [('password', '""'), ('products', '"both"')]:
        try:
            conn.execute(f'ALTER TABLE allowed_users ADD COLUMN {col} TEXT NOT NULL DEFAULT {default}')
        except sqlite3.OperationalError:
            pass
    conn.commit()
    conn.close()

init_db()


def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def verify_user(contact, password):
    """Returns the user row if valid, else None."""
    contact = contact.strip().lower()
    conn = get_db()
    user = conn.execute(
        'SELECT * FROM allowed_users WHERE LOWER(contact) = ? AND is_active = 1',
        (contact,)
    ).fetchone()
    conn.close()
    if user is None or user['password'] != password:
        return None
    return user


def user_has_product(product):
    """Check if the logged-in user has access to a specific product."""
    user_products = session.get('user_products', '')
    return user_products == 'both' or user_products == product


def login_required(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('user_logged_in'):
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('is_admin'):
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated

# ─────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────
MARGIN_LEFT = 45
MARGIN_RIGHT = 45
FOOTER_HEIGHT = 50
LINE_HEIGHT_TEXT = 17       # pixels per text line at font_size 11
LINE_HEIGHT_OPTION = 15     # pixels per option line at font_size 10
MATH_TAG_RE = _re.compile(r'(\[MATH\].*?\[/MATH\])', _re.DOTALL)


# ─────────────────────────────────────────────────────────
# KNOWLEDGE BASE
# ─────────────────────────────────────────────────────────
def read_knowledge_base():
    """Reads the Class 8 Math context file to use as RAG context."""
    kb_path = os.path.join(os.path.dirname(__file__), 'data', 'class8_math_kb.txt')
    try:
        with open(kb_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading KB: {e}")
        return "Class 8 Mathematics Syllabus Context."


# ─────────────────────────────────────────────────────────
# AI QUESTION GENERATION
# ─────────────────────────────────────────────────────────
def _call_groq_for_type(client, q_type, count, topics_str, difficulty_str, knowledge_context):
    """Makes a focused API call for a single question type. Returns a list of questions."""
    type_rules = {
        'MCQ': ('Multiple Choice', 'must have "options" list with EXACTLY 4 answer strings, and "type" set to "MCQ"'),
        'FIB': ('Fill in the Blank', 'question string must contain the blank as "____________", "options" must be [], "type" set to "FIB"'),
        'DESC': ('Descriptive / Long Answer', 'Must be a numerical word problem that requires calculation (NO DEFINITIONS). "options" must be [], "type" set to "DESC"'),
    }
    type_name, type_rule = type_rules[q_type]

    mcq_example = ''
    if q_type == 'MCQ':
        mcq_example = '''
EXAMPLE OUTPUT:
{"questions": [
  {"solution_verification": "sqrt(2) and pi are irrational. 3/4 is rational. Therefore, 3/4 is the correct answer.", "question": "Which of the following is a rational number?", "options": ["[MATH]sqrt(2)[/MATH]", "[MATH]sqrt(3)[/MATH]", "[MATH]3/4[/MATH]", "[MATH]pi[/MATH]"], "type": "MCQ", "difficulty": "Easy"},
  {"solution_verification": "2x + 5 = 13 => 2x = 8 => x = 4. Options include 3, 4, 5, 6.", "question": "Solve [MATH]2x + 5 = 13[/MATH]. What is x?", "options": ["[MATH]x = 3[/MATH]", "[MATH]x = 4[/MATH]", "[MATH]x = 5[/MATH]", "[MATH]x = 6[/MATH]"], "type": "MCQ", "difficulty": "Medium"}
]}'''
    elif q_type == 'FIB':
        mcq_example = '''
EXAMPLE OUTPUT:
{"questions": [
  {"solution_verification": "The square root of 144 is 12.", "question": "The square root of 144 is ____________.", "options": [], "type": "FIB", "difficulty": "Easy"},
  {"solution_verification": "2^8 means 2 multiplied by itself 8 times, which is 256.", "question": "[MATH]2^8[/MATH] = ____________.", "options": [], "type": "FIB", "difficulty": "Medium"}
]}'''
    elif q_type == 'DESC':
        mcq_example = '''
EXAMPLE OUTPUT:
{"questions": [
  {"solution_verification": "Area = 8*5 = 40. Perimeter = 2*(8+5) = 26.", "question": "A rectangle has dimensions [MATH]8 x 5[/MATH] cm. Find its area and perimeter.", "options": [], "type": "DESC", "difficulty": "Medium"},
  {"solution_verification": "Equation: 3x - 5 = 10 => 3x = 15 => x = 5.", "question": "Three times a number decreased by 5 equals 10. Find the number.", "options": [], "type": "DESC", "difficulty": "Easy"}
]}'''

    prompt = f"""You are a Class 8 Mathematics teacher at Arth Academy.

KNOWLEDGE BASE (use ONLY these concepts):
{knowledge_context}

TASK: Generate EXACTLY {count} {type_name} questions.
Topics: {topics_str} (STRICT ADHERENCE REQUIRED - DO NOT test topics that are not in this list)
Difficulty distribution: {difficulty_str}
Complexity Requirement: Ensure the level of questions exactly matches standard CBSE and ICSE board exams and strictly follows the requested difficulty level percentages.
Uniqueness Seed: {random.randint(10000, 99999)} - Use this to ensure entirely novel and unseen questions compared to typical runs.

RULES:
1. Return ONLY a valid JSON object: {{"questions": [...]}}
2. Generate EXACTLY {count} questions — no more, no fewer.
3. Each question: {type_rule}
4. "difficulty" must be: "Easy", "Medium", or "Tough"
5. NON-REPETITION and TOPIC MATCH: All {count} questions must test vastly different sub-concepts strictly bounded by the given Topics.
6. MATH & TEXT FORMATTING - CRITICAL RULES:
   - NEVER wrap an entire sentence or english words in [MATH] tags! Math tags are ONLY for equations/numbers.
     WRONG: [MATH]The area is 5 cm[/MATH]
     RIGHT: The area is [MATH]5[/MATH] cm.
   - ABSOLUTELY NO LATEX: Do not use \\frac, \\circ, \\pi, \\sqrt, etc. 
   - Use standard keyboard text: 1/4, 90 degrees, pi, sqrt(2).
   - Use × for multiplication, ÷ for division.
   - Use ^ for powers: 2^8, x^2.
   - NEVER use $ sign. Write currency as Rs. 120.
7. NO DIAGRAMS: Do NOT create any questions that refer to a figure or diagram (e.g., do not say "in the figure below"). These are strictly text-only questions.
8. NUMERICAL SOLVING ONLY: Do NOT ask for definitions or facts (e.g., "What is a polygon?", "State Euler's formula"). Every question must be a math problem that requires calculation or algebraic solving.
9. DOUBLE CHECK ACCURACY: You MUST internally verify your math. Include a "solution_verification" string field for EVERY question where you solve the problem step-by-step to logically ensure the question is valid, accurate, fully solvable, and perfectly tied to ONLY the requested Topics BEFORE writing the "question". 
{mcq_example}"""

    response = client.chat.completions.create(
        model='llama-3.3-70b-versatile',
        messages=[{'role': 'user', 'content': prompt}],
        response_format={"type": "json_object"},
        temperature=0.4,
        max_tokens=3000,
    )
    raw = response.choices[0].message.content.strip()
    parsed = json.loads(raw)
    if isinstance(parsed, dict):
        qs = next((v for v in parsed.values() if isinstance(v, list)), [])
    else:
        qs = parsed if isinstance(parsed, list) else []

    result = []
    for q in qs:
        if not isinstance(q, dict):
            continue
        opts = q.get('options', [])
        if not isinstance(opts, list):
            opts = []
        result.append({
            'question': auto_wrap_math(str(q.get('question', ''))),
            'options': [auto_wrap_math(str(o)) for o in opts],
            'type': q_type,
            'difficulty': str(q.get('difficulty', 'Medium')),
        })

    # Strict trim to requested count
    return result[:count]


def auto_wrap_math(text):
    """
    Post-process AI text to catch common math patterns written outside [MATH] tags:
    - sqrt(...) → [MATH]sqrt(...)[/MATH]
    - Numbers with ^ (e.g. 2^8) → [MATH]2^8[/MATH]
    Already-wrapped [MATH]...[/MATH] segments are left untouched.
    """
    # First, protect existing [MATH] blocks
    protected = []
    def protect(m):
        protected.append(m.group(0))
        return f'\x00{len(protected)-1}\x00'
    text = MATH_TAG_RE.sub(protect, text)

    # Wrap bare sqrt(...) that is not already inside [MATH]
    text = _re.sub(r'\bsqrt\(([^)]+)\)', r'[MATH]sqrt(\1)[/MATH]', text)
    # Wrap bare power expressions like 2^8, x^2 when sandwiched by spaces/punctuation
    text = _re.sub(r'(?<![\[/])(\w+\^\w+)(?![\]])', r'[MATH]\1[/MATH]', text)

    # Restore protected blocks
    for i, block in enumerate(protected):
        text = text.replace(f'\x00{i}\x00', block)
    return text


def _call_groq_for_geometry(client, count, topics_str, difficulty_str, knowledge_context):
    """Calls Groq to generate geometry diagram questions with structured shape data."""
    GEO_EXAMPLE = '''EXAMPLE OUTPUT:
{"questions": [
  {
    "solution_verification": "Triangle parameters: base=6, height=4. Area = 0.5 * 6 * 4 = 12. Mathematically solvable. Checking if it maps exactly to requested topics... yes.",
    "question": "Find the area of triangle ABC shown below.",
    "geometry": {"type": "triangle", "base": 6, "height": 4, "labels": ["A","B","C"], "dimensions": {"base": "6 cm", "height": "4 cm"}},
    "options": [],
    "type": "GEO",
    "difficulty": "Medium"
  }
]}'''

    prompt = f"""You are a Class 8 Mathematics teacher at Arth Academy.

KNOWLEDGE BASE:
{knowledge_context}

TASK: Generate EXACTLY {count} diagram-based geometry questions.
Complexity Requirement: Ensure the level of questions exactly matches standard CBSE and ICSE board exams and strictly follows the requested difficulty level percentages.
Topics: {topics_str} (STRICT ADHERENCE REQUIRED - DO NOT test topics that are not in this list)
Difficulty distribution: {difficulty_str}
Uniqueness Seed: {random.randint(10000, 99999)} - Use this to ensure entirely novel and unseen questions.

RULES:
1. Return ONLY valid JSON: {{"questions": [...]}}
2. Generate EXACTLY {count} questions — no more, no fewer.
3. EVERY question MUST include a "geometry" key with shape data. ONLY generate visual questions if they are relevant to the selected Topics (e.g. Mensuration, Coordinate Geometry).
4. NON-REPETITION: All {count} questions must use completely different shapes/contexts strictly pertaining to selected Topics.
5. Each question has these keys:
   - "question": string (e.g. "A cylindrical tank has radius 7m and height 10m. Find the cost of painting its curved surface at Rs. 50 per sq m.")
   - "geometry": object:
       - "type": one of ["triangle", "right_triangle", "rectangle", "square", "circle", "parallelogram", "cube", "cuboid", "cylinder", "coordinate_graph", "bar_chart", "pie_chart"]
       - triangle/right_triangle: "base", "height" (int), "labels" (3 strings), "dimensions" (dict of strings)
       - rectangle/square: "width", "height" (int), "labels" (4 strings), "dimensions" (dict of strings)
       - circle: "radius" (int), "dimensions" (dict of strings)
       - parallelogram: "base", "height" (int), "labels" (4 strings), "dimensions" (dict of strings)
       - cube/cuboid: "width", "height", "depth" (int), "dimensions" (dict of strings e.g. "length", "width", "height")
       - cylinder: "radius", "height" (int), "dimensions" (dict of strings e.g. "radius", "height")
       - coordinate_graph: "points" (array of objects: [{{"label": "A", "x": 2, "y": 3}}]), "lines" (array of coordinate arrays: [[[2,3], [4,5]]]), "polygons" (array of coordinate arrays: [[[0,0], [3,0], [0,4]]]), "x_range" (array: [-5, 5]), "y_range" (array: [-5, 5])
       - bar_chart: "categories" (array of strings: ["Math", "Sci"]), "values" (array of numbers: [40, 60]), "x_label" (string), "y_label" (string), "title" (string)
       - pie_chart: "categories" (array of strings: ["Math", "Sci"]), "values" (array of numbers: [30, 70]), "title" (string)
   - "options": [] (always empty)
   - "type": "GEO"
   - "difficulty": "Easy", "Medium", or "Tough"
6. COMPLEXITY RULES: For Medium and Tough questions, create complex multi-step problems. DO NOT ask basic direct formula questions.
7. Provide fully detailed dimensions in the "dimensions" dict so the diagram is clearly labelled.
8. Use realistic integer dimensions between 3 and 15.
9. MATH FORMATTING - CRITICAL RULES:
   - NEVER use LaTeX (no \\frac, \\circ, \\pi). Use "1/4", "90 degrees", "pi".
   - NEVER use $ sign. Write currency as Rs. 120.
   - DO NOT wrap English phrasing in [MATH] tags.
10. ACCURACY & CONSISTENCY: The numbers, dimensions, and shape type referenced in the `question` text MUST mathematically and logically match the values inside the `geometry` object PERFECTLY. The question must be mathematically accurate, fully solvable, and perfectly represented by the diagram.
11. TOPIC-SHAPE MAPPING (CRITICAL):
    - If Topics include "Coordinate Geometry", you MUST generate "coordinate_graph" type geometry.
    - If Topics include "Data Handling" or "Pie Charts & Bar Graphs", you MUST generate "bar_chart" or "pie_chart".
    - If Topics include "Mensuration" or "Solid Shapes", use shapes like cylinder, cube, rectangle, etc.
12. DOUBLE CHECK ACCURACY: You MUST internally verify your math. Include a "solution_verification" string field for EVERY question where you calculate the solution specifically using the values provided in your "geometry" object. You must ensure the numbers are possible (e.g., triangle inequality holds, dimensions map properly to the topic requested).
13. STRICT TYPES ONLY: You are absolutely FORBIDDEN from using any "type" not explicitly listed in the shape list. Do not invent shapes. If making a graph, use ONLY "coordinate_graph".

{GEO_EXAMPLE}"""

    response = client.chat.completions.create(
        model='llama-3.3-70b-versatile',
        messages=[{'role': 'user', 'content': prompt}],
        response_format={"type": "json_object"},
        temperature=0.4,
        max_tokens=3000,
    )
    raw = response.choices[0].message.content.strip()
    parsed = json.loads(raw)
    if isinstance(parsed, dict):
        qs = next((v for v in parsed.values() if isinstance(v, list)), [])
    else:
        qs = parsed if isinstance(parsed, list) else []

    result = []
    for q in qs:
        if not isinstance(q, dict):
            continue
        geo_data = q.get('geometry', None)
        if not isinstance(geo_data, dict):
            geo_data = None
        result.append({
            'question': auto_wrap_math(str(q.get('question', ''))),
            'options': [],
            'type': 'GEO',
            'difficulty': str(q.get('difficulty', 'Medium')),
            'geometry': geo_data,
        })
    return result[:count]


def generate_rag_questions(counts, topics, difficulty, groq_api_key, geo_count=0, request_id=None):
    """Uses Groq API (separate call per type + optional geometry) to guarantee exact question counts."""
    if request_id:
        progress_store[request_id] = {"status": "Reading Knowledge Base...", "progress": 10}
    knowledge_context = read_knowledge_base()
    topics_str = ", ".join(topics)
    diff_str = (f"Easy: {difficulty.get('easy', 33)}%, "
                f"Medium: {difficulty.get('medium', 34)}%, "
                f"Tough: {difficulty.get('tough', 33)}%")

    client = Groq(api_key=groq_api_key)
    all_questions = []
    type_map = [('MCQ', counts.get('mcq', 0)), ('FIB', counts.get('fib', 0)), ('DESC', counts.get('desc', 0))]
    if geo_count > 0:
        type_map.append(('GEO', geo_count))
    n_types = sum(1 for _, c in type_map if c > 0)
    progress_step = 55 // max(n_types, 1)
    current_progress = 15

    try:
        for q_type, count in type_map:
            if count == 0:
                continue
            if request_id:
                progress_store[request_id] = {
                    "status": f"Generating {count} {q_type} questions...",
                    "progress": current_progress
                }
            if q_type == 'GEO':
                questions = _call_groq_for_geometry(
                    client, count, topics_str, diff_str, knowledge_context
                )
            else:
                questions = _call_groq_for_type(
                    client, q_type, count, topics_str, diff_str, knowledge_context
                )
            all_questions.extend(questions)
            current_progress = min(current_progress + progress_step, 80)

        if request_id:
            progress_store[request_id] = {"status": "Parsing AI response...", "progress": 82}

        return all_questions

    except Exception as e:
        print(f"Groq Generation Error: {traceback.format_exc()}")
        raise Exception(f"Failed to generate questions: {str(e)}")


# ─────────────────────────────────────────────────────────
# GEOMETRY DIAGRAM RENDERER
# ─────────────────────────────────────────────────────────
def draw_geometry_diagram(geo):
    """
    Renders a geometry shape described by the `geo` dict and returns a PNG BytesIO buffer.
    Supported shapes: triangle, right_triangle, rectangle, square, circle, parallelogram.
    geo dict must have at least: {"type": "triangle", ...}
    Returns None on failure.
    """
    import numpy as np
    try:
        shape = geo.get('type', 'triangle').lower()
        fig, ax = plt.subplots(figsize=(3.5, 2.8))
        ax.set_aspect('equal')
        ax.axis('off')
        fig.patch.set_facecolor('white')
        ax.set_facecolor('white')

        # ---- Color palette ----
        FILL   = '#EEF0FF'  # light indigo fill
        STROKE = '#3320E6'  # indigo border
        LABEL  = '#1a1a3e'  # dark text
        DIM    = '#5533cc'  # dimension annotation color

        def style_polygon(patch):
            patch.set_facecolor(FILL)
            patch.set_edgecolor(STROKE)
            patch.set_linewidth(1.8)

        def draw_dim(x, y, text, rot=0, ha='center', va='center'):
            """Draws dimension text with a white background so it is highly readable."""
            ax.text(x, y, text, fontsize=9.5, fontweight='bold', color=DIM,
                    ha=ha, va=va, rotation=rot,
                    bbox=dict(facecolor='white', edgecolor='none', alpha=0.85, pad=1.5))

        if shape in ('triangle', 'right_triangle'):
            # Default scalene or right triangle
            if shape == 'right_triangle':
                pts = np.array([[0, 0], [geo.get('base', 6), 0],
                                [0, geo.get('height', 4)]], dtype=float)
            else:
                b = geo.get('base', 6)
                h = geo.get('height', 4)
                pts = np.array([[0, 0], [b, 0], [b/2, h]], dtype=float)

            poly = plt.Polygon(pts, closed=True)
            style_polygon(poly)
            ax.add_patch(poly)
            ax.autoscale()
            margin = 1.6  # increased margin for label cutoff
            ax.set_xlim(pts[:,0].min()-margin, pts[:,0].max()+margin)
            ax.set_ylim(pts[:,1].min()-margin, pts[:,1].max()+margin)

            # Vertex labels
            v_labels = geo.get('labels', ['A', 'B', 'C'])
            offsets = [(-0.35, -0.35), (0.25, -0.35), (0.0, 0.25)]
            for (px, py), lbl, (ox, oy) in zip(pts, v_labels, offsets):
                ax.text(px+ox, py+oy, lbl, fontsize=11, fontweight='bold', color=LABEL, ha='center')

            # Side dimensions
            dims = geo.get('dimensions', {})
            if 'base' in dims:
                mid_x = (pts[0][0] + pts[1][0]) / 2
                mid_y = pts[0][1]
                draw_dim(mid_x, mid_y - 0.55, dims['base'])
            if 'height' in dims:
                if shape == 'right_triangle':
                    draw_dim(pts[0][0] - 0.6, pts[2][1]/2, dims['height'], rot=90)
                else:
                    # draw a dashed height line
                    apex = pts[2]
                    base_mid_y = pts[0][1]
                    ax.plot([apex[0], apex[0]], [base_mid_y, apex[1]], '--', color='#aaa', linewidth=1)
                    draw_dim(apex[0], (base_mid_y + apex[1])/2, dims['height'], ha='center')
            if shape == 'right_triangle':
                # Right-angle mark
                sq = plt.Polygon([[0,0],[0.4,0],[0.4,0.4],[0,0.4],[0,0]], closed=True, fill=False, edgecolor=STROKE, linewidth=1)
                ax.add_patch(sq)
            # Side labels (A, B, C sides)
            side_dims = geo.get('sides', {})
            for k, v in side_dims.items():
                draw_dim(*ax.get_xlim(), v)

        elif shape in ('rectangle', 'square'):
            w = geo.get('width', 6)
            h = geo.get('height', w if shape == 'square' else 4)
            rect = plt.Polygon([[0,0],[w,0],[w,h],[0,h]], closed=True)
            style_polygon(rect)
            ax.add_patch(rect)
            ax.set_xlim(-1.4, w+1.4)
            ax.set_ylim(-1.4, h+1.4)

            # Dimension annotations
            dims = geo.get('dimensions', {})
            if 'width' in dims or 'side' in dims:
                draw_dim(w/2, 0, dims.get('width', dims.get('side')))
            if 'height' in dims or 'side' in dims:
                draw_dim(0, h/2, dims.get('height', dims.get('side')), rot=90)

            # Corner labels
            v_labels = geo.get('labels', ['A', 'B', 'C', 'D'])
            corners = [(0,0),(w,0),(w,h),(0,h)]
            c_offsets = [(-0.3,-0.3),(0.3,-0.3),(0.3,0.3),(-0.3,0.3)]
            for (cx,cy), lbl, (ox,oy) in zip(corners, v_labels, c_offsets):
                ax.text(cx+ox, cy+oy, lbl, fontsize=10, fontweight='bold', color=LABEL, ha='center')

        elif shape == 'circle':
            r = geo.get('radius', 4)
            circ = plt.Circle((0, 0), r, fill=True)
            circ.set_facecolor(FILL)
            circ.set_edgecolor(STROKE)
            circ.set_linewidth(1.8)
            ax.add_patch(circ)
            ax.set_xlim(-r-1.4, r+1.4)
            ax.set_ylim(-r-1.4, r+1.4)

            dims = geo.get('dimensions', {})
            r_lbl = dims.get('radius', f'r = {r}')
            d_lbl = dims.get('diameter', None)
            if d_lbl:
                ax.plot([-r, r], [0, 0], '-', color=STROKE, linewidth=1.2)
                draw_dim(0, 0, d_lbl)
            else:
                ax.plot([0, r], [0, 0], '-', color=STROKE, linewidth=1.2)
                draw_dim(r/2, 0, r_lbl)
            # Center dot
            ax.plot(0, 0, 'o', color=STROKE, markersize=3)
            ax.text(-0.1, -0.4, 'O', fontsize=9, color=LABEL)

        elif shape == 'parallelogram':
            b = geo.get('base', 6)
            h = geo.get('height', 4)
            slant = geo.get('slant', 1.5)
            pts = np.array([[0,0],[b,0],[b+slant,h],[slant,h]], dtype=float)
            poly = plt.Polygon(pts, closed=True)
            style_polygon(poly)
            ax.add_patch(poly)
            margin = 1.6
            ax.set_xlim(-margin, b+slant+margin)
            ax.set_ylim(-margin, h+margin)
            dims = geo.get('dimensions', {})
            if 'base' in dims:
                draw_dim(b/2, 0, dims['base'])
            if 'height' in dims:
                ax.plot([slant, slant], [0, h], '--', color='#aaa', linewidth=1)
                draw_dim(slant, h/2, dims['height'])
            v_labels = geo.get('labels', ['A', 'B', 'C', 'D'])
            offsets = [(-0.3,-0.35),(0.3,-0.35),(0.3,0.25),(-0.3,0.25)]
            for (px,py), lbl, (ox,oy) in zip(pts, v_labels, offsets):
                ax.text(px+ox, py+oy, lbl, fontsize=10, fontweight='bold', color=LABEL, ha='center')

        elif shape in ('cube', 'cuboid'):
            w = geo.get('width', 5)
            h = geo.get('height', w if shape == 'cube' else 3)
            d = geo.get('depth', w if shape == 'cube' else 2)
            # draw front face
            front = np.array([[0,0],[w,0],[w,h],[0,h]], dtype=float)
            # draw back face (isometric offset)
            off_x, off_y = d*0.5, d*0.5
            back = front + [off_x, off_y]
            
            # draw back edges & face
            poly_back = plt.Polygon(back, closed=True, facecolor=FILL, edgecolor=STROKE, linewidth=1)
            ax.add_patch(poly_back)
            
            # connecting lines
            for i in range(4):
                ax.plot([front[i,0], back[i,0]], [front[i,1], back[i,1]], color=STROKE, linewidth=1)
            
            # draw front face (drawn last so it overlays back face)
            poly_front = plt.Polygon(front, closed=True, facecolor='#ffffff', alpha=0.9, edgecolor=STROKE, linewidth=1.8)
            ax.add_patch(poly_front)

            ax.set_xlim(-1.6, w + off_x + 1.6)
            ax.set_ylim(-1.6, h + off_y + 1.6)

            dims = geo.get('dimensions', {})
            # label width (bottom of front face)
            if 'length' in dims or 'width' in dims or 'side' in dims:
                draw_dim(w/2, 0, dims.get('length', dims.get('width', dims.get('side'))))
            # label height (left edge of front face)
            if 'height' in dims or 'side' in dims:
                draw_dim(0, h/2, dims.get('height', dims.get('side')), rot=90)
            # label depth (bottom right skewed edge)
            if 'depth' in dims or 'width' in dims or 'side' in dims:
                draw_dim(w + off_x/2, off_y/2, dims.get('depth', dims.get('width', dims.get('side'))), rot=45)

        elif shape == 'cylinder':
            from matplotlib.patches import Ellipse
            r = geo.get('radius', 3)
            h = geo.get('height', 5)
            # We fake 3D by drawing a flat ellipse on top and an ellipse on bottom.
            # bottom full ellipse
            bottom_back = Ellipse((0,0), 2*r, r*0.6, fill=False, edgecolor=STROKE, linestyle='--', linewidth=1)
            bottom_front = Ellipse((0,0), 2*r, r*0.6, fill=True, facecolor=FILL, edgecolor=STROKE, linewidth=1.8)
            bottom_front.set_clip_box(ax.bbox)
            bottom_front.set_clip_path(plt.Rectangle((-r, -r), 2*r, r, transform=ax.transData)) # Only show bottom half
            
            # Body box
            body = plt.Rectangle((-r, 0), 2*r, h, facecolor=FILL, edgecolor='none')
            
            top = Ellipse((0,h), 2*r, r*0.6, fill=True, facecolor='#ffffff', alpha=0.7, edgecolor=STROKE, linewidth=1.8)
            
            ax.add_patch(bottom_back)
            ax.add_patch(body)
            ax.add_patch(bottom_front)
            ax.add_patch(top)
            
            # Side lines
            ax.plot([-r, -r], [0, h], color=STROKE, linewidth=1.8)
            ax.plot([r, r], [0, h], color=STROKE, linewidth=1.8)

            ax.set_xlim(-r-1.5, r+1.5)
            ax.set_ylim(-r-0.5, h+r+0.5)

            dims = geo.get('dimensions', {})
            if 'radius' in dims:
                ax.plot([0, r], [h, h], '-', color=STROKE, linewidth=1)
                ax.plot(0, h, 'o', color=STROKE, markersize=3)
                draw_dim(r/2, h, dims['radius'])
            if 'height' in dims:
                ax.plot([r+0.5, r+0.5], [0, h], '-', color='#aaa', linewidth=1)
                draw_dim(r+0.5, h/2, dims['height'])

        elif shape == 'coordinate_graph':
            x_range = geo.get('x_range', [-5, 5])
            if isinstance(x_range, dict): x_range = [x_range.get('min', -5), x_range.get('max', 5)]
            if not (isinstance(x_range, list) and len(x_range) == 2): x_range = [-5, 5]
            
            y_range = geo.get('y_range', [-5, 5])
            if isinstance(y_range, dict): y_range = [y_range.get('min', -5), y_range.get('max', 5)]
            if not (isinstance(y_range, list) and len(y_range) == 2): y_range = [-5, 5]
            
            # Make the plot a bit larger for cartesian graphs
            fig.set_size_inches(4.0, 4.0)
            ax.set_aspect('equal')
            ax.grid(True, linestyle='--', color='#e0e0e0', zorder=1)
            ax.axhline(0, color='black', linewidth=1.2, zorder=2)
            ax.axvline(0, color='black', linewidth=1.2, zorder=2)
            
            ax.set_xlim(x_range[0]-1, x_range[1]+1)
            ax.set_ylim(y_range[0]-1, y_range[1]+1)
            
            # draw ticks
            ax.set_xticks(range(int(x_range[0]), int(x_range[1]+1)))
            ax.set_yticks(range(int(y_range[0]), int(y_range[1]+1)))
            ax.tick_params(axis='both', labelsize=8)
            ax.axis('on') # override the default off
            for spine in ax.spines.values():
                spine.set_visible(False)
            
            # Add axis labels
            ax.text(x_range[1]+1.2, 0, 'x', fontsize=10, fontweight='bold', va='center')
            ax.text(0, y_range[1]+1.2, 'y', fontsize=10, fontweight='bold', ha='center')

            # Draw polygons if provided
            polygons = geo.get('polygons', [])
            if isinstance(polygons, list):
                for poly_coords in polygons:
                    if isinstance(poly_coords, list) and len(poly_coords) >= 3:
                        try:
                            # if LLM mistakenly gave [{"x":1, "y":2}] convert to [1, 2]
                            clean_coords = []
                            for p in poly_coords:
                                if isinstance(p, dict): clean_coords.append([p.get('x',0), p.get('y',0)])
                                elif isinstance(p, (list, tuple)) and len(p) >= 2: clean_coords.append(p[:2])
                            if len(clean_coords) >= 3:
                                pts = np.array(clean_coords, dtype=float)
                                poly = plt.Polygon(pts, closed=True, alpha=0.3, facecolor=FILL, edgecolor=STROKE, linewidth=1.5, zorder=1.5)
                                ax.add_patch(poly)
                        except Exception: pass

            # Draw lines if provided
            lines = geo.get('lines', [])
            if isinstance(lines, list):
                for line_coords in lines:
                    if isinstance(line_coords, list) and len(line_coords) >= 2:
                        try:
                            clean_coords = []
                            for p in line_coords:
                                if isinstance(p, dict): clean_coords.append([p.get('x',0), p.get('y',0)])
                                elif isinstance(p, (list, tuple)) and len(p) >= 2: clean_coords.append(p[:2])
                            xs = [p[0] for p in clean_coords]
                            ys = [p[1] for p in clean_coords]
                            ax.plot(xs, ys, '-', color=STROKE, linewidth=1.5, zorder=2.5)
                        except Exception: pass

            # Draw points if provided
            points = geo.get('points', [])
            for pt in points:
                x, y = pt.get('x', 0), pt.get('y', 0)
                label = pt.get('label', '')
                ax.plot(x, y, 'o', color=STROKE, markersize=5, zorder=3)
                if label:
                    ax.text(x+0.2, y+0.2, f'{label}({x},{y})', fontsize=9, fontweight='bold', color=LABEL, zorder=4,
                            bbox=dict(facecolor='white', edgecolor='none', alpha=0.8, pad=0.5))

        elif shape == 'bar_chart':
            ax.set_aspect('auto')  # Must be auto for bar charts
            categories = geo.get('categories', ['A', 'B', 'C'])
            values = geo.get('values', [10, 20, 15])
            if not isinstance(categories, list) or not isinstance(values, list) or len(categories) != len(values) or len(categories) == 0:
                categories, values = ['A', 'B', 'C'], [10, 20, 15]
            
            # Convert values to float if needed
            try: values = [float(v) for v in values]
            except: values = [10] * len(categories)
            
            fig.set_size_inches(4.5, 3.5)
            ax.axis('on')
            for spine in ax.spines.values():
                spine.set_visible(True)
                spine.set_color('#333333')
            ax.bar(categories, values, color=FILL, edgecolor=STROKE, linewidth=1.5, zorder=2)
            ax.set_xlabel(geo.get('x_label', ''), fontsize=9, fontweight='bold')
            ax.set_ylabel(geo.get('y_label', ''), fontsize=9, fontweight='bold')
            if geo.get('title'):
                ax.set_title(geo.get('title', ''), fontsize=10, fontweight='bold', pad=10)
            ax.tick_params(axis='both', labelsize=8)
            ax.grid(axis='y', linestyle='--', color='#e0e0e0', zorder=1)

        elif shape == 'pie_chart':
            ax.set_aspect('equal')
            categories = geo.get('categories', ['A', 'B', 'C'])
            values = geo.get('values', [30, 50, 20])
            if not isinstance(categories, list) or not isinstance(values, list) or len(categories) != len(values) or len(categories) == 0:
                categories, values = ['A', 'B', 'C'], [30, 50, 20]
            
            try: values = [float(v) for v in values]
            except: values = [30, 50, 20]
            if sum(values) <= 0: values = [1] * len(categories)

            fig.set_size_inches(4.0, 4.0)
            # Create a nice palette
            colors = ['#EEF0FF', '#D0D5FF', '#A0AAFF', '#7080FF', '#4055FF', '#2030EE', '#1010CC']
            wedges, texts, autotexts = ax.pie(values, labels=categories, autopct='%1.1f%%',
                                              startangle=90, colors=colors[:len(values)],
                                              wedgeprops=dict(edgecolor=STROKE, linewidth=1.5))
            for t in texts:
                t.set_fontsize(9)
                t.set_fontweight('bold')
            for at in autotexts:
                at.set_fontsize(8)
                at.set_color('#111111')
                at.set_fontweight('bold')
            if geo.get('title'):
                ax.set_title(geo.get('title', ''), fontsize=10, fontweight='bold', pad=10)

        else:  # Fallback: generic polygon from vertices
            verts = geo.get('vertices', [[0,0],[5,0],[2.5,4]])
            pts = np.array(verts, dtype=float)
            poly = plt.Polygon(pts, closed=True)
            style_polygon(poly)
            ax.add_patch(poly)
            ax.autoscale()

        plt.tight_layout(pad=0.3)
        buf = io_module.BytesIO()
        fig.savefig(buf, format='png', dpi=130, bbox_inches='tight',
                    pad_inches=0.1, facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        plt.close('all')
        print(f'[WARN] Geometry render failed: {traceback.format_exc()}')
        return None


# ─────────────────────────────────────────────────────────
# COUNT ENFORCEMENT
# ─────────────────────────────────────────────────────────
def enforce_counts(questions, counts):
    """
    After AI generation, sort questions into MCQ/FIB/DESC buckets
    and truncate each bucket to the requested count.
    This guarantees the PDF always has the exact sections requested.
    """
    target = {
        'MCQ': counts.get('mcq', 0),
        'FIB': counts.get('fib', 0),
        'DESC': counts.get('desc', 0),
    }
    buckets = {'MCQ': [], 'FIB': [], 'DESC': []}
    for q in questions:
        t = q.get('type', 'DESC').upper()
        if t in buckets:
            buckets[t].append(q)

    result = []
    for qtype in ('MCQ', 'FIB', 'DESC'):
        needed = target[qtype]
        got = buckets[qtype][:needed]   # Truncate to required count
        result.extend(got)
        if len(got) < needed:
            print(f"[WARN] Requested {needed} {qtype} but AI only returned {len(got)}")

    return result


# ─────────────────────────────────────────────────────────
# MATH RENDERING
# ─────────────────────────────────────────────────────────
def render_math_to_image(latex_expr, font_size=12, dpi=150):
    """
    Renders a LaTeX expression to a PNG byte buffer using matplotlib.
    Returns None on failure (caller must display plain text fallback).
    """
    try:
        # Use matplotlib's mathtext renderer — wrap in $...$
        expr_wrapped = f'${latex_expr}$'
        fig, ax = plt.subplots(figsize=(0.1, 0.1))
        ax.axis('off')
        txt = ax.text(
            0.5, 0.5, expr_wrapped,
            ha='center', va='center',
            fontsize=font_size,
            transform=ax.transAxes,
            usetex=False,
        )
        fig.canvas.draw()
        bbox = txt.get_window_extent(renderer=fig.canvas.get_renderer())
        padding = 0.06
        w_in = max((bbox.width / dpi) + padding, 0.05)
        h_in = max((bbox.height / dpi) + padding, 0.05)
        fig.set_size_inches(w_in, h_in)
        fig.patch.set_facecolor('white')
        ax.set_facecolor('white')

        buf = io_module.BytesIO()
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                    pad_inches=0.02, facecolor='white')
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        plt.close('all')
        return None


def sanitize_math_expr(expr):
    """
    Convert LaTeX backslash commands AND plain sqrt(...) to Unicode/clean equivalents
    so that matplotlib's mathtext renderer displays correct symbols.
    """
    substitutions = [
        (r'\\times', '×'),
        (r'\\cdot', '·'),
        (r'\\div', '÷'),
        (r'\\neq', '≠'),
        (r'\\leq', '≤'),
        (r'\\geq', '≥'),
        (r'\\approx', '≈'),
        (r'\\pm', '±'),
        (r'\\pi', 'π'),
        (r'\\infty', '∞'),
        (r'\\circ', '°'),
        # \frac{a}{b} → (a)/(b)
        (r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)'),
        # Fallback if parsed as form feed \f
        (r'\x0crac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)'),
        # \sqrt{n} → √n
        (r'\\sqrt\{([^}]*)\}', r'√(\1)'),
        (r'\\sqrt', '√'),
        # plain sqrt(n) → √n  (the AI may write this without backslash)
        (r'\bsqrt\(([^)]+)\)', r'√(\1)'),
        (r'\bsqrt\b', '√'),
        # \left and \right (invisible)
        (r'\\left', ''),
        (r'\\right', ''),
        # Generic: strip remaining backslash commands
        (r'\\([a-zA-Z]+)', r'\1'),
    ]
    result = expr
    for pattern, replacement in substitutions:
        result = _re.sub(pattern, replacement, result)
    return result


def _parse_segments(text):
    """
    Split text into a list of ('text', str) or ('math', latex_str) tuples.
    Uses [MATH]...[/MATH] as delimiters.
    """
    parts = MATH_TAG_RE.split(text)
    segments = []
    for part in parts:
        if part.startswith('[MATH]') and part.endswith('[/MATH]'):
            latex = part[6:-7].strip()
            if latex:
                segments.append(('math', latex))
        elif part:
            segments.append(('text', part))
    return segments


def _math_to_inline(latex_expr, font_size, max_h):
    """
    Returns (ImageReader, draw_w, draw_h) or None if rendering fails.
    """
    # Clean up LaTeX commands → Unicode before rendering
    clean_expr = sanitize_math_expr(latex_expr)
    buf = render_math_to_image(clean_expr, font_size=font_size + 1, dpi=150)
    if buf is None:
        return None
    try:
        ir = ImageReader(buf)
        iw, ih = ir.getSize()
        if ih == 0:
            return None
        scale = min(max_h / ih, 3.0)
        return ir, iw * scale, ih * scale
    except Exception:
        return None


# ─────────────────────────────────────────────────────────
# TEXT + MATH DRAWING ENGINE
# ─────────────────────────────────────────────────────────
def draw_rich_text(p, text, x, y, font_size=11, max_width=520, bold=False):
    """
    Draws text that may contain [MATH]...[/MATH] expressions.
    Handles word-wrapping and inline math image rendering.
    Returns the new y position after all lines are drawn.

    Strategy:
    - Collect tokens (words/spaces for text, images for math) into a line buffer
    - When adding a token would overflow max_width, flush and start a new line
    - Commit the line by drawing all items left-to-right at the current y
    """
    font_name = "Helvetica-Bold" if bold else "Helvetica"
    line_h = font_size + 6
    max_img_h = font_size + 8

    segments = _parse_segments(text)
    has_math = any(t == 'math' for t, _ in segments)

    # Pure text fast path — use ReportLab's own splitter for best results
    if not has_math:
        p.setFont(font_name, font_size)
        p.setFillColorRGB(0, 0, 0)
        # simpleSplit handles text-only wrapping perfectly
        lines = simpleSplit(text, font_name, font_size, max_width)
        for ln in lines:
            p.drawString(x, y, ln)
            y -= line_h
        return y - 4

    # Mixed path: build a line-buffer of (type, value, width, height)
    # type='text' -> value=str, type='math' -> value=ImageReader
    line_buf = []
    line_w = 0.0
    cursor_x = x

    def flush_line():
        nonlocal y, cursor_x, line_w
        tx = x
        for dtype, dval, dw, dh in line_buf:
            if dtype == 'text':
                p.setFont(font_name, font_size)
                p.setFillColorRGB(0, 0, 0)
                p.drawString(tx, y, dval)
            else:
                # Vertically center math image on the text baseline
                img_y = y - dh + font_size + 2
                p.drawImage(dval, tx, img_y, width=dw, height=dh)
            tx += dw
        y -= line_h
        cursor_x = x
        line_w = 0.0
        line_buf.clear()

    def add_token(dtype, dval, dw, dh):
        nonlocal cursor_x, line_w
        # If token alone exceeds max_width, force it on its own line
        if line_w > 0 and cursor_x + dw > x + max_width:
            flush_line()
        line_buf.append((dtype, dval, dw, dh))
        cursor_x += dw
        line_w += dw

    for seg_type, seg_val in segments:
        if not seg_val:
            continue
        if seg_type == 'text':
            # Split preserving spaces
            tokens = _re.split(r'(\s+)', seg_val)
            for token in tokens:
                if not token:
                    continue
                tw = p.stringWidth(token, font_name, font_size)
                add_token('text', token, tw, font_size)
        else:
            result = _math_to_inline(seg_val, font_size, max_img_h)
            if result:
                ir, dw, dh = result
                add_token('math', ir, dw, dh)
            else:
                # Fallback: render LaTeX as italic text in brackets
                fallback = f'[{seg_val}]'
                fw = p.stringWidth(fallback, "Helvetica-Oblique", font_size)
                p.setFont("Helvetica-Oblique", font_size)
                add_token('text', fallback, fw, font_size)

    if line_buf:
        flush_line()

    return y - 4  # extra paragraph gap


# ─────────────────────────────────────────────────────────
# PAGE ELEMENTS
# ─────────────────────────────────────────────────────────
def add_watermark(p, width, height):
    """Adds a diagonal 'Arth Academy' watermark to the current page."""
    p.saveState()
    p.setFont("Helvetica-Bold", 52)
    p.setFillColor(colors.Color(0.62, 0.60, 0.85, alpha=0.10))
    p.translate(width / 2, height / 2)
    p.rotate(45)
    p.drawCentredString(0, 0, "Arth Academy")
    p.restoreState()
    p.setFillColorRGB(0, 0, 0)
    p.setStrokeColorRGB(0, 0, 0)


def draw_page_header(p, width, height, topic, is_first_page=True, branding_opts=None):
    """Draws the branded header. Returns starting y for content."""
    if branding_opts is None:
        branding_opts = {}

    # Purple header banner
    header_h = 68
    p.setFillColorRGB(0.31, 0.27, 0.9)
    p.rect(0, height - header_h, width, header_h, fill=1, stroke=0)

    # Academy name
    p.setFillColorRGB(1, 1, 1)
    p.setFont("Helvetica-Bold", 22)
    p.drawString(MARGIN_LEFT, height - 42, "Arth Academy")

    # Right side info
    p.setFont("Helvetica", 10)
    max_marks = branding_opts.get('max_marks', '').strip() or "___ / ___"
    p.drawRightString(width - MARGIN_RIGHT, height - 26, "Class 8  ·  Mathematics")
    p.drawRightString(width - MARGIN_RIGHT, height - 44, f"Topic: {topic}")
    p.drawRightString(width - MARGIN_RIGHT, height - 60, f"Max Marks: {max_marks}")

    content_start_y = height - header_h - 14

    if is_first_page:
        # Student info bar
        p.setFillColorRGB(0.95, 0.95, 1.0)
        p.rect(MARGIN_LEFT - 5, content_start_y - 52, width - MARGIN_LEFT - MARGIN_RIGHT + 10, 52, fill=1, stroke=0)
        p.setFillColorRGB(0, 0, 0)
        p.setFont("Helvetica", 10.5)
        row1_y = content_start_y - 20
        row2_y = content_start_y - 40
        
        # dynamic fields
        s_name = branding_opts.get('student_name', '').strip() or "_______________________________"
        t_date = branding_opts.get('test_date', '').strip() or "_________________"
        c_sec  = branding_opts.get('class_sec', '').strip() or "___________"
        r_no   = branding_opts.get('roll_no', '').strip() or "___________"
        t_time = branding_opts.get('test_time', '').strip() or "___ Mins"
        if t_time and "Mins" not in t_time and t_time != "___ Mins":
            t_time += " Mins"

        p.drawString(MARGIN_LEFT, row1_y, f"Student Name: {s_name}")
        p.drawRightString(width - MARGIN_RIGHT, row1_y, f"Date: {t_date}")
        p.drawString(MARGIN_LEFT, row2_y, f"Class / Section: {c_sec}")
        p.drawString(260, row2_y, f"Roll No.: {r_no}")
        p.drawRightString(width - MARGIN_RIGHT, row2_y, f"Time: {t_time}")

        # Divider
        p.setStrokeColorRGB(0.31, 0.27, 0.9)
        p.setLineWidth(1.2)
        divider_y = content_start_y - 58
        p.line(MARGIN_LEFT, divider_y, width - MARGIN_RIGHT, divider_y)
        p.setFillColorRGB(0, 0, 0)
        p.setStrokeColorRGB(0, 0, 0)
        p.setLineWidth(1)

        # General instructions
        p.setFont("Helvetica-Oblique", 8.5)
        p.setFillColorRGB(0.35, 0.35, 0.35)
        instructions = ("Instructions: Answer all questions. Write clearly. "
                        "Marks are given as shown. No calculators unless specified.")
        p.drawString(MARGIN_LEFT, divider_y - 12, instructions)
        p.setFillColorRGB(0, 0, 0)

        return divider_y - 26  # y start for questions
    else:
        p.setFillColorRGB(0, 0, 0)
        return content_start_y - 14


def draw_page_footer(p, width, page_num, total_pages=None, branding_opts=None):
    """Draws footer with page number and optional branding contact info."""
    if branding_opts is None:
        branding_opts = {}
        
    footer_y = 30
    p.setFillColorRGB(0.85, 0.85, 0.95)
    p.rect(0, 0, width, footer_y + 12, fill=1, stroke=0)
    p.setStrokeColorRGB(0.31, 0.27, 0.9)
    p.setLineWidth(0.8)
    p.line(0, footer_y + 12, width, footer_y + 12)

    p.setFont("Helvetica", 8.5)
    p.setFillColorRGB(0.2, 0.2, 0.3)
    
    # Custom branding
    phone = branding_opts.get('phone_no', '').strip()
    email = branding_opts.get('email_id', '').strip()
    brand_parts = []
    if phone:
        brand_parts.append(f"Phone: {phone}")
    if email:
        brand_parts.append(f"Email: {email}")
    brand_text = "  |  ".join(brand_parts)
    if not brand_text:
        brand_text = "Class 8 Mathematics"
        
    p.drawString(MARGIN_LEFT, footer_y, brand_text)
    page_label = f"Page {page_num}" + (f" / {total_pages}" if total_pages else "")
    p.drawRightString(width - MARGIN_RIGHT, footer_y, page_label)
    p.setFillColorRGB(0, 0, 0)
    p.setStrokeColorRGB(0, 0, 0)


def draw_section_header(p, width, y, label):
    """Draws a colored section separator (e.g. 'Section A: MCQ'). Returns new y."""
    y -= 10
    p.setFillColorRGB(0.20, 0.17, 0.75)
    p.roundRect(MARGIN_LEFT - 5, y - 4, width - MARGIN_LEFT - MARGIN_RIGHT + 10, 20,
                4, fill=1, stroke=0)
    p.setFillColorRGB(1, 1, 1)
    p.setFont("Helvetica-Bold", 11)
    p.drawString(MARGIN_LEFT, y + 4, label)
    p.setFillColorRGB(0, 0, 0)
    return y - 22


def draw_difficulty_badge(p, y, difficulty):
    """Draws a small color-coded difficulty badge to the right of the question number y."""
    colors_map = {
        'Easy': (0.13, 0.68, 0.30),
        'Medium': (0.85, 0.55, 0.01),
        'Tough': (0.78, 0.13, 0.13),
    }
    c = colors_map.get(difficulty, (0.5, 0.5, 0.5))
    p.setFillColorRGB(*c)
    label = difficulty.upper()
    p.setFont("Helvetica-Bold", 7)
    badge_w = p.stringWidth(label, "Helvetica-Bold", 7) + 8
    p.roundRect(MARGIN_LEFT, y + 2, badge_w, 10, 2, fill=1, stroke=0)
    p.setFillColorRGB(1, 1, 1)
    p.drawString(MARGIN_LEFT + 4, y + 4, label)
    p.setFillColorRGB(0, 0, 0)


# ─────────────────────────────────────────────────────────
# MAIN PDF BUILDER
# ─────────────────────────────────────────────────────────
def create_branded_pdf(questions, topic, branding_opts=None):
    """Creates a professional Arth Academy branded PDF. Returns a BytesIO buffer."""
    buf = BytesIO()
    p = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    page_num = 1

    # Usable content width
    content_w = width - MARGIN_LEFT - MARGIN_RIGHT

    add_watermark(p, width, height)
    y = draw_page_header(p, width, height, topic, is_first_page=True, branding_opts=branding_opts)

    section_labels = {
        'MCQ': 'Section A: Multiple Choice Questions',
        'FIB': 'Section B: Fill in the Blanks',
        'DESC': 'Section C: Descriptive Questions',
        'GEO': 'Section D: Diagram-Based Questions',
    }
    current_section = None

    def ensure_space(needed, draw_section_hdr=False):
        """If not enough space remains, start a new page. Returns possibly-updated y."""
        nonlocal y, page_num
        extra = 32 if draw_section_hdr else 0
        if y - (needed + extra) < FOOTER_HEIGHT:
            draw_page_footer(p, width, page_num, branding_opts=branding_opts)
            p.showPage()
            page_num += 1
            add_watermark(p, width, height)
            y = draw_page_header(p, width, height, topic, is_first_page=False, branding_opts=branding_opts)

    for i, q in enumerate(questions, 1):
        q_type = q.get('type', 'DESC').upper()
        q_diff = q.get('difficulty', 'Medium')
        raw_question = q.get('question', '')
        options = q.get('options', [])
        if not isinstance(options, list):
            options = []

        # ── Conservative height estimate ──────────────────
        # Count approximate line count for question text
        q_chars = len(raw_question)
        est_q_lines = max(1, math.ceil(q_chars / 60))
        est_q_h = est_q_lines * LINE_HEIGHT_TEXT + 12

        # Options height
        if options:
            max_opt_chars = max(len(str(o)) for o in options)
            # If any option has math or is long, use vertical layout (all stacked)
            has_math_in_opts = any('[MATH]' in str(o) for o in options)
            use_grid = (not has_math_in_opts) and (max_opt_chars < 22)
            if use_grid:
                est_opt_h = (len(options) // 2 + 1) * LINE_HEIGHT_OPTION
            else:
                est_opt_h = len(options) * LINE_HEIGHT_OPTION * 2
        else:
            est_opt_h = 0

        # Geometry diagram height
        geo_h = 160 if q_type == 'GEO' else 0   # diagram image height estimate

        # Answer lines for FIB/DESC
        ans_h = 20 if q_type == 'FIB' else (50 if q_type in ('DESC', 'GEO') else 0)

        est_total = est_q_h + est_opt_h + ans_h + geo_h + 28  # 28 = inter-question gap

        # ── Section header ─────────────────────────────────
        need_section_hdr = (q_type != current_section) and (q_type in section_labels)
        ensure_space(est_total, draw_section_hdr=need_section_hdr)

        if need_section_hdr:
            current_section = q_type
            y = draw_section_header(p, width, y, section_labels[q_type])

        # ── Difficulty badge ────────────────────────────────
        draw_difficulty_badge(p, y, q_diff)
        y -= 14  # space below badge before question text

        # ── Question text ───────────────────────────────────
        question_text = f"Q{i}.  {raw_question}"
        p.setFont("Helvetica-Bold", 11)
        y = draw_rich_text(p, question_text, MARGIN_LEFT, y, font_size=11,
                           max_width=content_w, bold=True)

        # ── Geometry Diagram ────────────────────────────────
        if q_type == 'GEO':
            geo_data = q.get('geometry', None)
            if geo_data and isinstance(geo_data, dict):
                diagram_buf = draw_geometry_diagram(geo_data)
                if diagram_buf:
                    try:
                        ir = ImageReader(diagram_buf)
                        iw, ih = ir.getSize()
                        # Scale diagram to at most half content width, max 150pt high
                        max_diag_w = content_w * 0.55
                        max_diag_h = 150
                        scale = min(max_diag_w / iw, max_diag_h / ih, 1.0)
                        diag_w, diag_h = iw * scale, ih * scale
                        # Check space then draw centered
                        if y - diag_h < FOOTER_HEIGHT:
                            draw_page_footer(p, width, page_num, branding_opts=branding_opts)
                            p.showPage()
                            page_num += 1
                            add_watermark(p, width, height)
                            y = draw_page_header(p, width, height, topic, is_first_page=False, branding_opts=branding_opts)
                        diag_x = MARGIN_LEFT + (content_w - diag_w) / 2  # center it
                        # Light border around diagram
                        p.setStrokeColorRGB(0.75, 0.75, 0.9)
                        p.setLineWidth(0.8)
                        p.roundRect(diag_x - 4, y - diag_h - 4, diag_w + 8, diag_h + 8,
                                    4, fill=0, stroke=1)
                        p.drawImage(ir, diag_x, y - diag_h, width=diag_w, height=diag_h)
                        y -= diag_h + 12
                        p.setStrokeColorRGB(0, 0, 0)
                    except Exception as ex:
                        print(f'[WARN] Could not embed diagram: {ex}')
                else:
                    # Fallback placeholder box
                    y -= 8
                    p.setStrokeColorRGB(0.7, 0.7, 0.85)
                    p.setLineWidth(0.7)
                    p.roundRect(MARGIN_LEFT, y - 90, content_w, 90, 4, fill=0, stroke=1)
                    p.setFont("Helvetica-Oblique", 9)
                    p.setFillColorRGB(0.5, 0.5, 0.6)
                    p.drawCentredString(MARGIN_LEFT + content_w/2, y - 48, "[Diagram placeholder]")
                    y -= 98
                    p.setFillColorRGB(0, 0, 0)
                    p.setStrokeColorRGB(0, 0, 0)

        # ── MCQ Options ─────────────────────────────────────
        if options:
            labels = ['(A)', '(B)', '(C)', '(D)']
            has_math_in_opts = any('[MATH]' in str(o) for o in options)
            max_opt_chars = max(len(str(o)) for o in options)
            use_grid = (not has_math_in_opts) and (max_opt_chars < 22)

            if use_grid:
                # 2-column grid
                col_w = content_w / 2
                for row in range(0, len(options), 2):
                    for col in range(2):
                        idx = row + col
                        if idx < len(options):
                            opt_label = labels[idx]
                            opt_txt = f"{opt_label}  {options[idx]}"
                            ox = MARGIN_LEFT + col * col_w + (10 if col else 0)
                            p.setFont("Helvetica", 10)
                            p.setFillColorRGB(0, 0, 0)
                            p.drawString(ox, y, opt_txt)
                    y -= LINE_HEIGHT_OPTION + 2
            else:
                # Vertical — each option full-width (handles long text and math)
                for idx, opt in enumerate(options):
                    opt_txt = f"{labels[idx]}  {opt}"
                    p.setFont("Helvetica", 10)
                    y = draw_rich_text(p, opt_txt, MARGIN_LEFT + 8, y,
                                       font_size=10, max_width=content_w - 10, bold=False)

        # ── Answer lines for FIB / DESC ─────────────────────
        if q_type == 'FIB':
            y -= 4
            p.setStrokeColorRGB(0.55, 0.55, 0.55)
            p.setLineWidth(0.6)
            p.line(MARGIN_LEFT + 10, y, width - MARGIN_RIGHT, y)
            y -= 6
        # Answer lines for GEO / DESC
        if q_type in ('GEO', 'DESC'):
            p.setStrokeColorRGB(0.7, 0.7, 0.7)
            p.setLineWidth(0.5)
            for _ in range(3):
                y -= 18
                p.line(MARGIN_LEFT + 10, y, width - MARGIN_RIGHT, y)
            y -= 6

        y -= 18  # inter-question spacing

    # Final page
    draw_page_footer(p, width, page_num, branding_opts=branding_opts)
    p.showPage()
    p.save()
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────
# FLASK ROUTES
# ─────────────────────────────────────────────────────────
@app.route('/')
def home():
    if not session.get('user_logged_in'):
        return redirect(url_for('login_page'))
    return redirect(url_for('dashboard'))


@app.route('/dashboard')
@login_required
def dashboard():
    products = session.get('user_products', 'both')
    return render_template('dashboard.html', products=products)


@app.route('/generator')
@login_required
def generator_home():
    if not user_has_product('generator'):
        return render_template('no_access.html', product='Test Generator'), 403
    return render_template('generator/index.html')


# ─── TEST SERIES ROUTES ─────────────────────────────────
# In-memory store for uploaded test sessions
test_sessions = {}


def _to_standard_question(q, test_id=None):
    """Convert question to scalable structure: subject, chapter, question_text, question_type, options with is_correct."""
    q_type = str(q.get('type', 'mcq')).strip().lower()
    question_type = 'MCQ' if q_type == 'mcq' else 'fill_in_the_blanks'
    opts_raw = q.get('options', [])
    answer = str(q.get('answer', '')).strip()

    options = []
    if q_type == 'mcq':
        for o in opts_raw:
            if isinstance(o, dict):
                options.append({'text': str(o.get('text', '')), 'is_correct': bool(o.get('is_correct'))})
            else:
                opt_text = str(o).strip()
                options.append({'text': opt_text, 'is_correct': opt_text.lower() == answer.lower()})

    correct_display = answer
    if options:
        for o in options:
            if o.get('is_correct'):
                correct_display = o['text']
                break

    out = {
        'id': q.get('id'),
        'type': q_type,
        'subject': str(q.get('subject', '')).strip(),
        'chapter': str(q.get('chapter', '')).strip(),
        'question_text': str(q.get('question_text', q.get('question', ''))).strip(),
        'question': str(q.get('question_text', q.get('question', ''))).strip(),
        'question_type': question_type,
        'difficulty': str(q.get('difficulty', '')).strip(),
        'marks': int(q.get('marks', 0)) if q.get('marks') else 0,
        'options': options,
        'answer': answer,
        'correct_display': correct_display,
        'has_image': bool(q.get('has_image')),
        'image_count': int(q.get('image_count', 0)),
    }
    if test_id and out['has_image']:
        out['image_url'] = f'/series/image/{test_id}/{out["id"]}'
    return out


def _parse_diagram_string(raw):
    """Parse a diagram column value into a geometry dict.
    Accepts JSON string or simple syntax like: triangle:base=6,height=4,labels=A;B;C
    Auto-generates 'dimensions' dict for the renderer to display labels.
    """
    if not raw or str(raw).strip() == '':
        return None
    raw = str(raw).strip()
    if raw.startswith('{'):
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            return None
    if ':' not in raw:
        return None
    shape_type, _, params_str = raw.partition(':')
    geo = {'type': shape_type.strip().lower()}
    for part in params_str.split(','):
        part = part.strip()
        if '=' not in part:
            continue
        k, _, v = part.partition('=')
        k, v = k.strip(), v.strip()
        if ';' in v:
            geo[k] = [s.strip() for s in v.split(';')]
        else:
            try:
                geo[k] = int(v)
            except ValueError:
                try:
                    geo[k] = float(v)
                except ValueError:
                    geo[k] = v

    # Auto-generate dimensions dict for the renderer
    dims = {}
    shape = geo.get('type', '')
    if shape in ('triangle', 'right_triangle'):
        if 'base' in geo:
            dims['base'] = f"{geo['base']} cm"
        if 'height' in geo:
            dims['height'] = f"{geo['height']} cm"
    elif shape in ('rectangle', 'square'):
        if 'width' in geo:
            dims['width'] = f"{geo['width']} cm"
        if 'height' in geo:
            dims['height'] = f"{geo['height']} cm"
        if shape == 'square' and 'width' in geo:
            dims['side'] = f"{geo['width']} cm"
    elif shape == 'circle':
        if 'radius' in geo:
            dims['radius'] = f"r = {geo['radius']} cm"
    elif shape == 'parallelogram':
        if 'base' in geo:
            dims['base'] = f"{geo['base']} cm"
        if 'height' in geo:
            dims['height'] = f"{geo['height']} cm"
    elif shape in ('cube', 'cuboid'):
        if 'width' in geo:
            dims['length'] = f"{geo['width']} cm"
        if 'height' in geo:
            dims['height'] = f"{geo['height']} cm"
        if 'depth' in geo:
            dims['depth'] = f"{geo['depth']} cm"
        if shape == 'cube' and 'width' in geo:
            dims['side'] = f"{geo['width']} cm"
    elif shape == 'cylinder':
        if 'radius' in geo:
            dims['radius'] = f"r = {geo['radius']} cm"
        if 'height' in geo:
            dims['height'] = f"h = {geo['height']} cm"
    if dims:
        geo['dimensions'] = dims
    return geo


def parse_questions_from_excel(file_stream):
    """Parse questions and extract embedded images from an Excel file."""
    import openpyxl
    from zipfile import ZipFile
    from xml.etree import ElementTree as ET

    raw_bytes = file_stream.read()
    file_stream.seek(0)
    df = pd.read_excel(BytesIO(raw_bytes))
    df.columns = [c.strip().lower() for c in df.columns]
    df = df.fillna('')

    # embedded_images: row -> list of image bytes (multiple images per row)
    embedded_images = {}
    orphan_images = []

    try:
        zf = ZipFile(BytesIO(raw_bytes))
        media_files = {}
        for n in zf.namelist():
            if n.startswith('xl/media/'):
                media_files[n] = zf.read(n)
                media_files[os.path.basename(n)] = zf.read(n)

        rels_map = {}
        for rname in zf.namelist():
            if 'drawings/_rels/' in rname and rname.endswith('.rels'):
                rel_tree = ET.fromstring(zf.read(rname))
                for rel in rel_tree:
                    rid = rel.get('Id')
                    target_raw = rel.get('Target', '')
                    if not target_raw:
                        continue
                    target_clean = target_raw.replace('../', '').lstrip('/')
                    full = 'xl/' + target_clean if not target_clean.startswith('xl/') else target_clean
                    rels_map[rid] = full
                    rels_map[rid + '_alt'] = target_clean

        ns_xdr = 'http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing'
        ns_a = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

        for dname in zf.namelist():
            if 'drawings/drawing' not in dname or not dname.endswith('.xml'):
                continue
            tree = ET.fromstring(zf.read(dname))
            for anchor_tag in ['twoCellAnchor', 'oneCellAnchor', 'absoluteAnchor']:
                for anchor in tree.findall(f'.//{{{ns_xdr}}}{anchor_tag}'):
                    from_el = anchor.find(f'{{{ns_xdr}}}from')
                    row = None
                    if from_el is not None:
                        row_el = from_el.find(f'{{{ns_xdr}}}row')
                        if row_el is not None and row_el.text is not None:
                            row = int(row_el.text)

                    blip = anchor.find(f'.//{{{ns_a}}}blip')
                    if blip is None:
                        continue
                    embed_id = blip.get(f'{{{ns_r}}}embed')
                    if not embed_id:
                        continue

                    path = rels_map.get(embed_id) or rels_map.get(embed_id + '_alt')
                    img_data = media_files.get(path) or media_files.get(os.path.basename(path or ''))
                    if img_data:
                        if row is not None:
                            embedded_images.setdefault(row, []).append(img_data)
                        else:
                            orphan_images.append(img_data)

        zf.close()
    except Exception as e:
        print(f"[WARN] Excel ZIP image extraction failed: {e}")

    if not embedded_images and not orphan_images:
        try:
            wb = openpyxl.load_workbook(BytesIO(raw_bytes))
            ws = wb.active
            for img in getattr(ws, '_images', []):
                try:
                    anchor = getattr(img, 'anchor', None)
                    row = None
                    if anchor and hasattr(anchor, '_from'):
                        row = anchor._from.row
                    elif anchor and hasattr(anchor, 'anchorFrom'):
                        row = anchor.anchorFrom.row
                    data = None
                    if hasattr(img, '_data') and callable(img._data):
                        data = img._data()
                    elif hasattr(img, 'ref') and hasattr(img.ref, 'read'):
                        data = img.ref.read()
                    if data and row is not None:
                        embedded_images.setdefault(row, []).append(data)
                except Exception:
                    continue
            wb.close()
        except Exception as e2:
            print(f"[WARN] Excel openpyxl fallback failed: {e2}")

    questions = []
    images_store = {}
    # Excel: row 1 = header, row 2 = first data. DataFrame index 0 = first data row.
    # Try multiple row mappings: i+2 (Excel row 2 for df 0), i+1, i+3, i, i+4
    for i, row in df.iterrows():
        q_id = int(row.get('id', i + 1))
        q_type = str(row.get('type', 'mcq')).strip().lower()
        opts = str(row.get('options', ''))
        opts_list = [o.strip() for o in opts.split('|')] if opts else []
        q = {
            'id': q_id,
            'subject': str(row.get('subject', '')).strip(),
            'chapter': str(row.get('chapter', '')).strip(),
            'type': q_type,
            'question': str(row.get('question', '')),
            'answer': str(row.get('answer', '')),
            'options': opts_list,
            'difficulty': str(row.get('difficulty', '')).strip(),
            'marks': row.get('marks', 0),
            'has_image': False,
            'image_count': 0,
        }
        try:
            q['marks'] = int(q['marks']) if q['marks'] else 0
        except (ValueError, TypeError):
            q['marks'] = 0

        imgs = None
        for candidate_row in [i + 2, i + 1, i + 3, i, i + 4, i + 5]:
            if candidate_row in embedded_images:
                imgs = embedded_images[candidate_row]
                break
        if imgs:
            images_store[q_id] = imgs if isinstance(imgs, list) else [imgs]
            q['has_image'] = True
            q['image_count'] = len(images_store[q_id])
        questions.append(q)

    if orphan_images and not any(q.get('has_image') for q in questions):
        for idx, q in enumerate(questions):
            if idx < len(orphan_images):
                images_store[q['id']] = [orphan_images[idx]]
                q['has_image'] = True
                q['image_count'] = 1

    return questions, images_store


def _extract_docx_images_via_zip(file_stream):
    """Extract images from docx - maps block_index -> list of image bytes. Uses same block order as _iter_docx_blocks_in_order."""
    from zipfile import ZipFile
    from xml.etree import ElementTree as ET
    raw = file_stream.read()
    file_stream.seek(0)
    block_images = {}
    ns_r = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
    try:
        zf = ZipFile(BytesIO(raw))
        if 'word/document.xml' not in zf.namelist():
            zf.close()
            return block_images

        rels = {}
        for n in zf.namelist():
            if n == 'word/_rels/document.xml.rels':
                tree = ET.fromstring(zf.read(n))
                for rel in tree:
                    rid = rel.get('Id')
                    target = (rel.get('Target') or '').strip()
                    if not target:
                        continue
                    target = target.replace('../', '').lstrip('/')
                    if 'media/' in target or target.endswith(('.png', '.jpg', '.jpeg', '.emf', '.wmf', '.gif')):
                        path = 'word/' + target if not target.startswith('word/') else target
                        rels[rid] = path
                break

        media = {}
        for n in zf.namelist():
            if n.startswith('word/media/'):
                data = zf.read(n)
                media[n] = data
                media[os.path.basename(n)] = data

        doc_tree = ET.fromstring(zf.read('word/document.xml'))
        ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        body = doc_tree.find(f'{{{ns_w}}}body') or doc_tree.find('.//{*}body') or doc_tree
        if body is None:
            body = doc_tree

        def _get_imgs_from_elem(elem):
            imgs, seen = [], set()
            for el in elem.iter():
                tag = (el.tag or '').split('}')[-1] if '}' in (el.tag or '') else ''
                rid = None
                if tag == 'blip':
                    rid = el.get(ns_r + 'embed')
                elif 'imagedata' in tag:
                    rid = el.get(ns_r + 'id') or el.get(ns_r + 'href')
                if rid and rid in rels and rid not in seen:
                    path = rels[rid]
                    data = media.get(path) or media.get(os.path.basename(path))
                    if data:
                        imgs.append(data)
                        seen.add(rid)
            return imgs

        block_idx = 0
        for child in list(body):
            tag = (child.tag or '').split('}')[-1] if '}' in (child.tag or '') else ''
            if tag == 'p':
                imgs = _get_imgs_from_elem(child)
                if imgs:
                    block_images[block_idx] = imgs
                block_idx += 1
            elif tag == 'tbl':
                for tr in child.findall('.//{*}tr'):
                    for tc in tr.findall('.//{*}tc'):
                        for p in tc.findall('.//{*}p'):
                            imgs = _get_imgs_from_elem(p)
                            if imgs:
                                block_images[block_idx] = imgs
                            block_idx += 1
        if not block_images and media:
            all_imgs = []
            seen = set()
            for el in doc_tree.iter():
                tag = (el.tag or '').split('}')[-1] if '}' in (el.tag or '') else ''
                rid = None
                if tag == 'blip':
                    rid = el.get(ns_r + 'embed')
                elif 'imagedata' in tag:
                    rid = el.get(ns_r + 'id') or el.get(ns_r + 'href')
                if rid and rid in rels and rid not in seen:
                    path = rels[rid]
                    data = media.get(path) or media.get(os.path.basename(path))
                    if data:
                        all_imgs.append(data)
                        seen.add(rid)
            if all_imgs:
                block_images['_orphans'] = all_imgs
        zf.close()
    except Exception as e:
        print(f"[WARN] docx ZIP image extraction failed: {e}")
    return block_images


def _extract_docx_paragraph_images(para, rels, zip_images=None, block_idx=None):
    """Extract images from a docx paragraph. Uses zip_images if provided (more reliable)."""
    if zip_images is not None and block_idx is not None and block_idx in zip_images:
        return zip_images[block_idx]

    seen_ids = set()
    images = []

    def _add_blob_for_rid(rid):
        if not rid or rid in seen_ids:
            return
        try:
            if hasattr(rels, 'get') and rid in rels:
                blob = rels[rid].target_part.blob
            elif hasattr(rels, 'related_parts') and rid in rels.related_parts:
                blob = rels.related_parts[rid].blob
            else:
                return
            if blob:
                seen_ids.add(rid)
                images.append(blob)
        except Exception:
            pass

    ns_a = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
    ns_r = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
    ns_v = '{urn:schemas-microsoft-com:vml}'
    ns_o = '{urn:schemas-microsoft-com:office:office}'
    el = para._element

    for blip in el.iter(f'{ns_a}blip'):
        _add_blob_for_rid(blip.get(f'{ns_r}embed'))
    for imgdata in el.iter(f'{ns_v}imagedata'):
        _add_blob_for_rid(imgdata.get(f'{ns_r}id'))
        _add_blob_for_rid(imgdata.get(f'{ns_r}href'))
    for ole in el.iter(f'{ns_o}OLEObject'):
        _add_blob_for_rid(ole.get(f'{ns_r}id'))
    for shape in el.iter(f'{ns_v}shape'):
        for child in shape:
            if 'imagedata' in child.tag:
                _add_blob_for_rid(child.get(f'{ns_r}id'))

    return images


def _iter_docx_blocks_in_order(doc):
    """Yield paragraphs and table-cell paragraphs in document order (body + tables)."""
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para


def parse_questions_from_docx(file_stream):
    from docx import Document
    raw = file_stream.read()
    file_stream.seek(0)
    doc = Document(BytesIO(raw))
    rels = doc.part.rels
    zip_images = _extract_docx_images_via_zip(BytesIO(raw))
    questions = []
    images_store = {}
    q_id = 0
    current_q = None
    block_idx = 0

    try:
        block_iter = _iter_docx_blocks_in_order(doc)
    except Exception:
        block_iter = doc.paragraphs

    for para in block_iter:
        text = (para.text or '').strip()

        para_images = _extract_docx_paragraph_images(para, rels, zip_images=zip_images, block_idx=block_idx)
        block_idx += 1

        # Is this the start of a new question?
        is_new_q = bool(text) and text.lower().startswith('q') and ('.' in text[:5] or ')' in text[:5])

        if is_new_q:
            if current_q:
                current_q['image_count'] = len(images_store.get(current_q['id'], []))
                questions.append(current_q)
            q_id += 1
            q_text = _re.sub(r'^[Qq]\d+[\.\)]\s*', '', text).strip()
            current_q = {
                'id': q_id, 'type': 'fill_in_the_blanks',
                'subject': '', 'chapter': '', 'difficulty': '', 'marks': 0,
                'question': q_text, 'options': [], 'answer': '',
                'has_image': False, 'image_count': 0,
            }

        # Rule: everything below (options, diagrams, images) belongs to current question until next Q
        if para_images and current_q:
            images_store.setdefault(current_q['id'], []).extend(para_images)
            current_q['has_image'] = True

        if not text or is_new_q:
            continue

        if text.lower().startswith('answer:') or text.lower().startswith('ans:'):
            if current_q:
                ans = _re.sub(r'^(?:answer|ans)\s*:\s*', '', text, flags=_re.IGNORECASE).strip()
                current_q['answer'] = ans

        elif text.lower().startswith(('a)', 'b)', 'c)', 'd)', 'a.', 'b.', 'c.', 'd.')):
            if current_q:
                current_q['type'] = 'mcq'
                opt_text = _re.sub(r'^[a-dA-D][\.\)]\s*', '', text).strip()
                current_q['options'].append(opt_text)

        elif current_q and not current_q['question']:
            current_q['question'] = text

    if current_q:
        current_q['image_count'] = len(images_store.get(current_q['id'], []))
        questions.append(current_q)

    orphans = zip_images.get('_orphans', []) if isinstance(zip_images, dict) else []
    if orphans and not any(q.get('has_image') for q in questions):
        for idx, q in enumerate(questions):
            if idx < len(orphans):
                images_store[q['id']] = [orphans[idx]]
                q['has_image'] = True
                q['image_count'] = 1

    return questions, images_store


def generate_sample_excel():
    data = [
        {"id": 1, "type": "mcq", "subject": "Mathematics", "chapter": "Rational Numbers", "question": "Which of the following is a rational number?", "options": "√2|π|0|√3", "answer": "0", "difficulty": "Easy", "marks": 1, "diagram": ""},
        {"id": 2, "type": "mcq", "subject": "Mathematics", "chapter": "Squares", "question": "What is the square of 15?", "options": "225|255|125|325", "answer": "225", "difficulty": "Easy", "marks": 1, "diagram": ""},
        {"id": 3, "type": "mcq", "subject": "Mathematics", "chapter": "Linear Equations", "question": "If 3x + 5 = 20, what is the value of x?", "options": "3|5|15|4", "answer": "5", "difficulty": "Medium", "marks": 2, "diagram": ""},
        {"id": 4, "type": "fill_in_the_blanks", "subject": "Mathematics", "chapter": "Cubes", "question": "The cube root of 512 is ____.", "options": "", "answer": "8", "difficulty": "Easy", "marks": 1, "diagram": ""},
        {"id": 5, "type": "fill_in_the_blanks", "subject": "Mathematics", "chapter": "Mensuration", "question": "Find the area of the triangle shown below.", "options": "", "answer": "12", "difficulty": "Medium", "marks": 2, "diagram": ""},
        {"id": 6, "type": "mcq", "subject": "Mathematics", "chapter": "Mensuration", "question": "What is the area of the circle shown below?", "options": "12.56|15.70|28.27|50.27", "answer": "28.27", "difficulty": "Medium", "marks": 4, "diagram": ""},
        {"id": 7, "type": "fill_in_the_blanks", "subject": "Mathematics", "chapter": "Mensuration", "question": "Find the perimeter of the rectangle shown below.", "options": "", "answer": "22", "difficulty": "Medium", "marks": 2, "diagram": ""},
    ]
    df = pd.DataFrame(data)
    buf = BytesIO()
    df.to_excel(buf, index=False)
    buf.seek(0)
    return buf


def _make_blue_circle_png():
    """Create a simple blue circle PNG."""
    try:
        from PIL import Image, ImageDraw
        size = 120
        img = Image.new('RGBA', (size, size), (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.ellipse((4, 4, size - 4, size - 4), fill=(100, 149, 237, 255), outline=(70, 130, 180, 255))
        buf = BytesIO()
        img.save(buf, format='PNG')
        buf.seek(0)
        return buf
    except Exception:
        return None


def generate_sample_docx():
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    doc.add_heading('Sample Test Questions', level=1)
    doc.add_paragraph('Format: Start each question with Q1. Q2. etc.\n'
                      'Add options as A) B) C) D) on separate lines.\n'
                      'Write Answer: on a separate line.\n'
                      'For diagrams, paste or draw your image below the question—everything belongs to that question until the next Q.\n')
    doc.add_paragraph('Q1. Which of the following is a rational number?')
    doc.add_paragraph('A) √2')
    doc.add_paragraph('B) π')
    doc.add_paragraph('C) 0')
    doc.add_paragraph('D) √3')
    doc.add_paragraph('Answer: 0')
    doc.add_paragraph('')
    doc.add_paragraph('Q2. What is the square of 15?')
    doc.add_paragraph('A) 225')
    doc.add_paragraph('B) 255')
    doc.add_paragraph('C) 125')
    doc.add_paragraph('D) 325')
    doc.add_paragraph('Answer: 225')
    doc.add_paragraph('')
    doc.add_paragraph('Q3. The cube root of 512 is ____.')
    doc.add_paragraph('Answer: 8')
    doc.add_paragraph('')
    doc.add_paragraph('Q4. Find the area of the triangle shown below.')
    doc.add_paragraph('[Paste or draw your diagram here]')
    doc.add_paragraph('Answer: 12')
    doc.add_paragraph('')
    doc.add_paragraph('Q5. What is the area of the circle shown below?')
    p_img = doc.add_paragraph()
    circle_buf = _make_blue_circle_png()
    if circle_buf:
        run = p_img.add_run()
        run.add_picture(circle_buf, width=Inches(1.5))
    doc.add_paragraph('A) 12.56')
    doc.add_paragraph('B) 15.70')
    doc.add_paragraph('C) 28.27')
    doc.add_paragraph('D) 50.27')
    doc.add_paragraph('Answer: 28.27')
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route('/series')
@login_required
def series_home():
    if not user_has_product('series'):
        return render_template('no_access.html', product='Test Series'), 403
    return render_template('series/setup.html')


@app.route('/series/sample/excel')
@login_required
def series_sample_excel():
    buf = generate_sample_excel()
    return send_file(buf, as_attachment=True, download_name='sample_questions.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/series/sample/word')
@login_required
def series_sample_word():
    buf = generate_sample_docx()
    return send_file(buf, as_attachment=True, download_name='sample_questions.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/series/start', methods=['POST'])
@login_required
def series_start():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files['file']
    filename = file.filename.lower()
    timer_mins = int(request.form.get('timer', 10))
    total_marks = int(request.form.get('marks', 0))

    try:
        if filename.endswith('.xlsx') or filename.endswith('.xls'):
            questions, images = parse_questions_from_excel(file)
        elif filename.endswith('.docx'):
            questions, images = parse_questions_from_docx(file)
        else:
            return jsonify({"error": "Unsupported file format. Upload .xlsx or .docx"}), 400
    except Exception as e:
        return jsonify({"error": f"Failed to parse file: {str(e)}"}), 400

    if not questions:
        return jsonify({"error": "No questions found in the file. Check the format."}), 400

    test_id = str(uuid.uuid4())
    normalized = [_to_standard_question(q, test_id=None) for q in questions]
    test_sessions[test_id] = {
        'questions': normalized,
        'timer': timer_mins,
        'marks': total_marks,
        'images': images,
    }
    for q in test_sessions[test_id]['questions']:
        if q.get('has_image'):
            q['image_url'] = f'/series/image/{test_id}/{q["id"]}'
    return jsonify({"test_id": test_id})


def _detect_image_mimetype(data):
    if data[:3] == b'\xff\xd8\xff':
        return 'image/jpeg'
    if data[:4] == b'GIF8':
        return 'image/gif'
    if data[:4] == b'RIFF' and len(data) > 11 and data[8:12] == b'WEBP':
        return 'image/webp'
    if data[:2] == b'BM':
        return 'image/bmp'
    if len(data) >= 44 and data[40:44] == b' EMF':
        return 'image/x-emf'
    if len(data) >= 4 and data[:4] == b'\x01\x00\x00\x00':
        return 'image/x-emf'
    if len(data) >= 4 and data[:4] == b'\xd7\xcd\xc6\x9a':
        return 'image/x-wmf'
    return 'image/png'


def _convert_emf_wmf_to_png(img_bytes):
    """Convert EMF/WMF (Word pasted images) to PNG using ImageMagick. Returns PNG bytes or None."""
    import subprocess
    import tempfile
    ext = '.emf' if _detect_image_mimetype(img_bytes) == 'image/x-emf' else '.wmf'
    try:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f_in:
            f_in.write(img_bytes)
            path_in = f_in.name
        path_out = path_in + '.png'
        try:
            for cmd in [['convert', path_in, path_out], ['magick', path_in, path_out]]:
                r = subprocess.run(cmd, capture_output=True, timeout=15)
                if r.returncode == 0 and os.path.exists(path_out):
                    with open(path_out, 'rb') as f:
                        return f.read()
        finally:
            for p in [path_in, path_out]:
                try:
                    os.unlink(p)
                except OSError:
                    pass
    except Exception:
        pass
    return None


@app.route('/series/image/<test_id>/<int:q_id>')
@app.route('/series/image/<test_id>/<int:q_id>/<int:idx>')
@login_required
def series_image(test_id, q_id, idx=0):
    """Serve an embedded image extracted from an uploaded file."""
    ts = test_sessions.get(test_id)
    if not ts:
        return "Not found", 404
    img_list = ts.get('images', {}).get(q_id)
    if not img_list:
        return "No image", 404
    if isinstance(img_list, bytes):
        img_list = [img_list]
    if idx < 0 or idx >= len(img_list):
        return "Image index out of range", 404
    img_bytes = img_list[idx]
    mimetype = _detect_image_mimetype(img_bytes)
    if mimetype in ('image/x-emf', 'image/x-wmf'):
        png_bytes = _convert_emf_wmf_to_png(img_bytes)
        if png_bytes:
            img_bytes = png_bytes
            mimetype = 'image/png'
        else:
            try:
                from PIL import Image
                buf = BytesIO(img_bytes)
                img = Image.open(buf)
                out = BytesIO()
                img.save(out, format='PNG')
                img_bytes = out.getvalue()
                mimetype = 'image/png'
            except Exception:
                try:
                    from PIL import Image
                    img = Image.new('RGB', (200, 80), color=(240, 240, 240))
                    from PIL import ImageDraw
                    d = ImageDraw.Draw(img)
                    d.text((10, 25), 'Pasted image: use Paste as Picture', fill=(100, 100, 100))
                    out = BytesIO()
                    img.save(out, format='PNG')
                    img_bytes = out.getvalue()
                    mimetype = 'image/png'
                except Exception:
                    pass
    buf = BytesIO(img_bytes)
    return send_file(buf, mimetype=mimetype)


@app.route('/series/diagram/<test_id>/<int:q_id>')
@login_required
def series_diagram(test_id, q_id):
    """Render and serve a geometry diagram as PNG for a specific question."""
    ts = test_sessions.get(test_id)
    if not ts:
        return "Not found", 404
    for q in ts['questions']:
        if q['id'] == q_id and q.get('geometry'):
            buf = draw_geometry_diagram(q['geometry'])
            if buf:
                return send_file(buf, mimetype='image/png')
    return "No diagram", 404


@app.route('/series/test/<test_id>')
@login_required
def series_test(test_id):
    ts = test_sessions.get(test_id)
    if not ts:
        return "Test not found or expired.", 404
    return render_template('series/index.html',
                           questions=ts['questions'],
                           timer=ts['timer'],
                           marks=ts['marks'],
                           test_id=test_id)


@app.route('/series/submit/<test_id>', methods=['POST'])
@login_required
def series_submit(test_id):
    ts = test_sessions.get(test_id)
    if not ts:
        return jsonify({"error": "Test session expired."}), 404
    try:
        questions_data = ts['questions']
        user_answers = request.get_json()
        score = 0
        total_questions = len(questions_data)
        results = []
        for q in questions_data:
            q_id = str(q['id'])
            user_ans = str(user_answers.get(q_id, "")).strip().lower()
            correct_display = str(q.get('correct_display', q.get('answer', ''))).strip()
            correct_ans = str(q.get('answer', '')).strip().lower()
            is_correct = user_ans == correct_ans or user_ans == correct_display.lower()
            if is_correct:
                score += 1
            results.append({
                "id": q['id'],
                "is_correct": is_correct,
                "correct_answer": correct_display or q.get('answer', ''),
                "user_answer": user_answers.get(q_id, "")
            })
        marks_per_q = ts['marks'] / total_questions if ts['marks'] and total_questions else 0
        total_scored = round(score * marks_per_q, 1) if ts['marks'] else score
        return jsonify({
            "score": score,
            "total": total_questions,
            "percentage": round((score / total_questions) * 100, 2) if total_questions else 0,
            "results": results,
            "marks_scored": total_scored,
            "marks_total": ts['marks']
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@app.route('/series/download_test/<test_id>')
@login_required
def series_download(test_id):
    ts = test_sessions.get(test_id)
    if not ts:
        return "Test not found.", 404
    questions_data = ts['questions']
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    p.setFont("Helvetica-Bold", 18)
    p.drawString(200, height - 50, "Arth Academy - Math Test")
    p.setFont("Helvetica", 12)
    p.drawString(50, height - 80, "Name: ________________________")
    p.drawString(350, height - 80, "Date: _____________")
    if ts['marks']:
        p.drawString(50, height - 100, f"Total Marks: {ts['marks']}")
        p.drawString(350, height - 100, f"Time: {ts['timer']} minutes")
        y = height - 130
    else:
        y = height - 110
    for i, q in enumerate(questions_data, 1):
        if y < 100:
            p.showPage()
            y = height - 50
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y, f"Q{i}: {q.get('question_text', q.get('question', ''))}")
        y -= 25
        p.setFont("Helvetica", 12)
        if str(q.get('type') or q.get('question_type', '')).lower() == 'mcq':
            labels = ["A)", "B)", "C)", "D)"]
            opt_x = 70
            for j in range(min(len(q.get('options', [])), 4)):
                opt = q['options'][j]
                opt_text = opt.get('text', opt) if isinstance(opt, dict) else opt
                p.drawString(opt_x, y, f"{labels[j]} {opt_text}")
                opt_x += 120
            y -= 30
        else:
            y -= 40
    p.showPage()
    p.save()
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='Arth_Academy_Test.pdf', mimetype='application/pdf')


@app.route('/login', methods=['GET', 'POST'])
def login_page():
    if request.method == 'POST':
        contact = request.form.get('contact', '').strip()
        password = request.form.get('password', '').strip()
        if not contact or not password:
            return render_template('login.html', error='Please enter both email/phone and password.')
        user = verify_user(contact, password)
        if user:
            session['user_logged_in'] = True
            session['user_contact'] = contact
            session['user_products'] = user['products'] or 'both'
            return redirect(url_for('dashboard'))
        else:
            return render_template('login.html', error='Invalid credentials or access denied. Contact admin.')
    return render_template('login.html')


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login_page'))


@app.route('/admin', methods=['GET', 'POST'])
def admin_login():
    if session.get('is_admin'):
        return redirect(url_for('admin_panel'))
    if request.method == 'POST':
        password = request.form.get('password', '')
        if password == ADMIN_PASSWORD:
            session['is_admin'] = True
            return redirect(url_for('admin_panel'))
        return render_template('admin_login.html', error='Wrong password.')
    return render_template('admin_login.html')


@app.route('/admin/panel')
@admin_required
def admin_panel():
    show_archived = request.args.get('archived', '0') == '1'
    conn = get_db()
    if show_archived:
        users = conn.execute('SELECT * FROM allowed_users ORDER BY added_on DESC').fetchall()
    else:
        users = conn.execute('SELECT * FROM allowed_users WHERE is_active >= 0 ORDER BY added_on DESC').fetchall()
    conn.close()
    return render_template('admin_panel.html', users=users, show_archived=show_archived)


@app.route('/admin/add', methods=['POST'])
@admin_required
def admin_add_user():
    contact = request.form.get('contact', '').strip().lower()
    name = request.form.get('name', '').strip()
    password = request.form.get('password', '').strip()
    products = request.form.get('products', 'both').strip()
    if not contact or not password:
        return redirect(url_for('admin_panel'))
    conn = get_db()
    try:
        conn.execute('INSERT INTO allowed_users (contact, name, password, products) VALUES (?, ?, ?, ?)', (contact, name, password, products))
        conn.commit()
    except sqlite3.IntegrityError:
        if not products:
            products = 'both'
        conn.execute('UPDATE allowed_users SET is_active = 1, name = ?, password = ?, products = ? WHERE LOWER(contact) = ?', (name, password, products, contact))
        conn.commit()
    conn.close()
    return redirect(url_for('admin_panel'))


@app.route('/admin/remove/<int:user_id>')
@admin_required
def admin_remove_user(user_id):
    conn = get_db()
    conn.execute('UPDATE allowed_users SET is_active = 0 WHERE id = ?', (user_id,))
    conn.commit()
    conn.close()
    return redirect(url_for('admin_panel'))


@app.route('/admin/activate/<int:user_id>')
@admin_required
def admin_activate_user(user_id):
    conn = get_db()
    conn.execute('UPDATE allowed_users SET is_active = 1 WHERE id = ?', (user_id,))
    conn.commit()
    conn.close()
    return redirect(url_for('admin_panel'))


@app.route('/admin/delete/<int:user_id>')
@admin_required
def admin_delete_user(user_id):
    conn = get_db()
    conn.execute('UPDATE allowed_users SET is_active = -1 WHERE id = ?', (user_id,))
    conn.commit()
    conn.close()
    return redirect(url_for('admin_panel'))


@app.route('/admin/logout')
def admin_logout():
    session.pop('is_admin', None)
    return redirect(url_for('admin_login'))


@app.route('/progress/<request_id>')
def progress(request_id):
    def generate():
        while True:
            if request_id in progress_store:
                data = progress_store[request_id]
                yield f"data: {json.dumps(data)}\n\n"
                if data.get("progress") == 100 or data.get("error"):
                    break
            else:
                yield f"data: {json.dumps({'error': 'Invalid Request ID'})}\n\n"
                break
            time.sleep(0.5)
    return Response(generate(), mimetype='text/event-stream')


def background_generation_task(request_id, counts, topics, difficulty, groq_api_key, geo_count=0, branding_opts=None):
    try:
        questions = generate_rag_questions(counts, topics, difficulty, groq_api_key,
                                           geo_count=geo_count, request_id=request_id)
        progress_store[request_id] = {"status": "Building PDF layout...", "progress": 90}
        title_topic = topics[0] if len(topics) == 1 else "Mixed Topics"
        pdf_buffer = create_branded_pdf(questions, title_topic, branding_opts=branding_opts)
        result_store[request_id] = {
            "buffer": pdf_buffer,
            "filename": f'Arth_Academy_Test_{title_topic.replace(" ", "_")}.pdf'
        }
        progress_store[request_id] = {"status": "Complete!", "progress": 100}
    except Exception as e:
        progress_store[request_id] = {"status": "Error", "progress": 0, "error": str(e)}


@app.route('/generate', methods=['POST'])
@login_required
def generate_route():
    data = request.json
    counts = {
        'mcq': int(data.get('mcq_count', 5)),
        'fib': int(data.get('fib_count', 0)),
        'desc': int(data.get('desc_count', 0)),
    }
    difficulty = {
        'easy': int(data.get('easy_pct', 33)),
        'medium': int(data.get('medium_pct', 34)),
        'tough': int(data.get('tough_pct', 33)),
    }
    branding_opts = {
        'student_name': data.get('student_name', ''),
        'class_sec': data.get('class_sec', ''),
        'roll_no': data.get('roll_no', ''),
        'test_date': data.get('test_date', ''),
        'test_time': data.get('test_time', ''),
        'max_marks': data.get('max_marks', ''),
        'phone_no': data.get('phone_no', ''),
        'email_id': data.get('email_id', ''),
    }
    geo_count = int(data.get('geo_count', 0))
    topics = data.get('topics', ['General Mathematics'])
    groq_api_key = os.environ.get('GROQ_API_KEY', '')

    if not groq_api_key:
        return jsonify({"error": "Server API key not configured. Contact admin."}), 500

    request_id = str(uuid.uuid4())
    progress_store[request_id] = {"status": "Initializing...", "progress": 5}

    thread = threading.Thread(
        target=background_generation_task,
        args=(request_id, counts, topics, difficulty, groq_api_key, geo_count, branding_opts)
    )
    thread.daemon = True
    thread.start()

    return jsonify({"request_id": request_id})


@app.route('/download/<request_id>')
@login_required
def download(request_id):
    if request_id not in result_store:
        return "PDF not found or expired.", 404
    result = result_store[request_id]
    return send_file(
        result["buffer"],
        as_attachment=True,
        download_name=result["filename"],
        mimetype='application/pdf'
    )


if __name__ == '__main__':
    app.run(debug=True, port=5002)
